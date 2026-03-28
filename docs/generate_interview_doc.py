"""Generate a clean, readable interview-ready .docx for the Text-to-SQL project."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Global style ──

style = doc.styles["Normal"]
style.font.name = "Segoe UI"
style.font.size = Pt(10.5)
style.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)
style.paragraph_format.line_spacing = 1.3

BLUE = RGBColor(0x1A, 0x56, 0xDB)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x55, 0x55, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = "DCE6F1"
HEADER_BG = "1A56DB"
ALT_ROW = "F0F4FA"
ACCENT_BG = "EBF0FF"
GREEN_BG = "E8F5E9"
ORANGE_BG = "FFF3E0"

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Segoe UI"
    hs.font.color.rgb = DARK
    hs.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    hs.paragraph_format.space_after = Pt(6)
    if level == 1:
        hs.font.size = Pt(18)
    elif level == 2:
        hs.font.size = Pt(14)
    else:
        hs.font.size = Pt(12)


def shade_cell(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def nice_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell in t.rows[0].cells:
        shade_cell(cell, HEADER_BG)

    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        r.font.name = "Segoe UI"

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            if ri % 2 == 1:
                shade_cell(cell, ALT_ROW)
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
            r.font.name = "Segoe UI"
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def section(title, level=1):
    h = doc.add_heading(title, level=level)
    if level == 1:
        p = doc.add_paragraph()
        r = p.add_run()
        r.font.size = Pt(1)
    return h


def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def bold_body(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10.5)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix + "  ")
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def colored_box(text, bg_color=ACCENT_BG):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    shade_cell(cell, bg_color)
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Consolas"
    r.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    doc.add_paragraph()


def qa(question, answer):
    p = doc.add_paragraph()
    r = p.add_run("Q: " + question)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE
    a = doc.add_paragraph("A: " + answer)
    a.paragraph_format.space_after = Pt(10)


# ═══════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════

doc.add_paragraph().space_before = Pt(80)

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_para.add_run("Text-to-SQL Platform")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = BLUE
r.font.name = "Segoe UI"

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Interview Documentation and Architecture Guide")
r.font.size = Pt(14)
r.font.color.rgb = GRAY
r.font.name = "Segoe UI"

doc.add_paragraph()

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = author.add_run("Rajat Joshi")
r.font.size = Pt(13)
r.font.color.rgb = DARK
r.bold = True

doc.add_page_break()

# ═══════════════════════════════════════
# 1. PROJECT OVERVIEW
# ═══════════════════════════════════════

section("1. Project Overview")

body("This project lets people ask questions in plain English and get back SQL queries. It uses a fine-tuned AI model to understand the question and the database structure, then writes the SQL automatically.")

body("For example, a user can paste their database tables and ask \"How many employees earn more than 50000?\" and the system will return: SELECT COUNT(*) FROM employees WHERE salary > 50000;")

section("What the system does", 2)
bullet("Takes a database schema (table definitions) and a question in English")
bullet("Generates an SQL query using a fine-tuned language model")
bullet("Returns a confidence score so users know how reliable the answer is")
bullet("Falls back to smart rule-based SQL if the AI model is not available")
bullet("Shows everything in a clean web dashboard with history tracking")

section("Tech Stack", 2)
nice_table(
    ["Layer", "Technology", "What it does"],
    [
        ["AI / ML", "PyTorch + HuggingFace + LoRA", "Trains and runs the SQL generation model"],
        ["Dataset", "Spider CSV (8,035 examples)", "Question and SQL pairs for training"],
        ["Backend", "FastAPI + Python", "REST API that receives questions and returns SQL"],
        ["Frontend", "Next.js + React + TypeScript + Tailwind", "Web interface for users to interact with"],
    ],
)

section("Key Numbers", 2)
nice_table(
    ["Metric", "Value", "What Does This Mean?"],
    [
        ["Training Loss", "0.40",
         "How wrong the model is on training data. Lower is better. 0.40 means the model learned the patterns well."],
        ["Eval Loss", "0.43",
         "How wrong the model is on new, unseen data. Close to training loss (0.40) means the model is not overfitting."],
        ["Token Accuracy", "87.5%",
         "Out of all the words/tokens the model generates, 87.5% match the correct answer exactly."],
        ["Model Adapter Size", "4.5 MB",
         "The fine-tuned part is tiny (4.5 MB) compared to the full model (2.2 GB). Easy to store and share."],
        ["Trainable Parameters", "0.10% (1.1M of 1.1B)",
         "Only 1 in every 1000 parameters gets updated. The rest stay frozen from the pretrained model."],
        ["Training Time", "~2.5 hours",
         "Trained on a MacBook with Apple Silicon. No expensive cloud GPU needed."],
        ["Inference Speed", "Under 1 second",
         "Time from question to SQL output. Fast enough for interactive use in the dashboard."],
        ["Dataset Size", "8,035 examples",
         "Number of English question + SQL answer pairs the model learned from (Spider dataset)."],
    ],
    col_widths=[3.5, 3, 9],
)

doc.add_page_break()

# ═══════════════════════════════════════
# 1B. HOW TO RUN + EXAMPLES
# ═══════════════════════════════════════

section("1B. How to Run the Application")

section("Quick Start (One Command)", 2)
colored_box(
    "  STEP 1: Clone the repository\n"
    "  $ git clone <repo-url>\n"
    "  $ cd text-to-sql-finetuning\n"
    "\n"
    "  STEP 2: Run everything with one command\n"
    "  $ ./run.sh\n"
    "\n"
    "  This script will:\n"
    "    - Create Python virtual environment\n"
    "    - Install all Python dependencies\n"
    "    - Install frontend npm packages\n"
    "    - Start the FastAPI backend on port 8000\n"
    "    - Start the Next.js frontend on port 3000\n"
    "\n"
    "  STEP 3: Open your browser\n"
    "  Go to: http://localhost:3000",
    GREEN_BG,
)

section("Manual Start (If You Prefer)", 2)
colored_box(
    "  Terminal 1 (Backend):\n"
    "  $ cd text-to-sql-finetuning\n"
    "  $ source venv/bin/activate\n"
    "  $ python api/serve.py\n"
    "  --> Backend running at http://localhost:8000\n"
    "\n"
    "  Terminal 2 (Frontend):\n"
    "  $ cd text-to-sql-finetuning/frontend\n"
    "  $ npm install\n"
    "  $ npm run dev\n"
    "  --> Frontend running at http://localhost:3000",
    ACCENT_BG,
)

section("Input and Output: What to Expect", 2)

body("The application takes two inputs and gives three outputs:")

nice_table(
    ["Direction", "Field", "What Is It", "Example"],
    [
        ["INPUT", "Database Schema", "The CREATE TABLE statements that describe your database",
         "CREATE TABLE employees (id INT, name TEXT, salary INT, dept TEXT);"],
        ["INPUT", "English Question", "A question about the data in plain English",
         "How many employees earn more than 50000?"],
        ["OUTPUT", "SQL Query", "The generated SQL that answers the question",
         "SELECT COUNT(*) FROM employees WHERE salary > 50000;"],
        ["OUTPUT", "Confidence", "A percentage showing how sure the model is (0-100%)",
         "92% (shown as green badge)"],
        ["OUTPUT", "Latency", "How long the system took to generate the answer",
         "340 ms"],
    ],
    col_widths=[2, 3, 5, 6],
)

section("Example 1: Counting with a Filter", 2)

colored_box(
    "  SCHEMA (what you paste):\n"
    "  CREATE TABLE employees (\n"
    "      id INT PRIMARY KEY,\n"
    "      name TEXT,\n"
    "      salary INT,\n"
    "      department TEXT\n"
    "  );\n"
    "\n"
    "  QUESTION (what you type):\n"
    "  How many employees earn more than 50000?\n"
    "\n"
    "  OUTPUT (what the system returns):\n"
    "  +----------------------------------------------+\n"
    "  | SQL:        SELECT COUNT(*)                  |\n"
    "  |             FROM employees                   |\n"
    "  |             WHERE salary > 50000;            |\n"
    "  |                                              |\n"
    "  | Confidence: 92%  [GREEN BADGE]               |\n"
    "  | Latency:    340 ms                           |\n"
    "  +----------------------------------------------+",
    GREEN_BG,
)

section("Example 2: Aggregation with Grouping", 2)

colored_box(
    "  SCHEMA:\n"
    "  CREATE TABLE orders (\n"
    "      id INT, customer_id INT, amount DECIMAL,\n"
    "      order_date DATE, status TEXT\n"
    "  );\n"
    "\n"
    "  QUESTION:\n"
    "  Show total sales amount for each status.\n"
    "\n"
    "  OUTPUT:\n"
    "  +----------------------------------------------+\n"
    "  | SQL:        SELECT status, SUM(amount)       |\n"
    "  |             FROM orders                      |\n"
    "  |             GROUP BY status;                 |\n"
    "  |                                              |\n"
    "  | Confidence: 88%  [GREEN BADGE]               |\n"
    "  | Latency:    410 ms                           |\n"
    "  +----------------------------------------------+",
    ACCENT_BG,
)

section("Example 3: Simple Select", 2)

colored_box(
    "  SCHEMA:\n"
    "  CREATE TABLE students (\n"
    "      id INT, name TEXT, major TEXT, gpa DECIMAL\n"
    "  );\n"
    "\n"
    "  QUESTION:\n"
    "  List all students in Computer Science.\n"
    "\n"
    "  OUTPUT:\n"
    "  +----------------------------------------------+\n"
    "  | SQL:        SELECT *                         |\n"
    "  |             FROM students                    |\n"
    "  |             WHERE major = 'Computer Science';|\n"
    "  |                                              |\n"
    "  | Confidence: 85%  [GREEN BADGE]               |\n"
    "  | Latency:    290 ms                           |\n"
    "  +----------------------------------------------+",
    GREEN_BG,
)

section("Example 4: Average Calculation", 2)

colored_box(
    "  SCHEMA:\n"
    "  CREATE TABLE products (\n"
    "      id INT, name TEXT, price DECIMAL,\n"
    "      category TEXT, in_stock BOOLEAN\n"
    "  );\n"
    "\n"
    "  QUESTION:\n"
    "  What is the average price of products in each category?\n"
    "\n"
    "  OUTPUT:\n"
    "  +----------------------------------------------+\n"
    "  | SQL:        SELECT category, AVG(price)      |\n"
    "  |             FROM products                    |\n"
    "  |             GROUP BY category;               |\n"
    "  |                                              |\n"
    "  | Confidence: 90%  [GREEN BADGE]               |\n"
    "  | Latency:    375 ms                           |\n"
    "  +----------------------------------------------+",
    ACCENT_BG,
)

doc.add_page_break()

# ═══════════════════════════════════════
# 2. WHAT I BUILT
# ═══════════════════════════════════════

section("2. What I Built (If Interviewer Asks: What Did You Implement?)")

body("I built every part of this project from scratch. Here is exactly what I implemented, broken down by area.")

section("AI and Machine Learning", 2)
numbered("Built the data pipeline that reads Spider CSV data and converts it into training-ready instruction prompts")
numbered("Wrote the training script that loads TinyLlama model, attaches LoRA adapters, and fine-tunes using SFTTrainer")
numbered("Created a configuration system using YAML files so all settings (model, training, LoRA) are easy to change")
numbered("Built an evaluation system that checks if generated SQL matches expected SQL, with breakdown by query type")
numbered("Implemented confidence scoring that uses the model's own token probabilities to tell users how sure it is")

section("Backend API", 2)
numbered("Built the FastAPI server with two endpoints: /health (status check) and /generate_sql (main feature)")
numbered("Added input validation so bad requests get clear error messages")
numbered("Implemented two paths: real AI model inference when available, and smart rule-based fallback when model is not loaded")
numbered("Built schema parsing that understands CREATE TABLE statements and extracts table/column information")

section("Frontend Web App", 2)
numbered("Built three pages: landing page, SQL generation dashboard, and a chat-style interface")
numbered("Created custom React hooks for API calls, health monitoring, and query history")
numbered("Built the SQL output display with color-coded confidence badges, latency info, and copy button")
numbered("Added dark/light theme, toast notifications, and responsive design")

section("DevOps", 2)
numbered("Created a single run.sh script that sets up the entire project and starts both servers")
numbered("Made the project installable as a Python package with command-line tools")

section("How the Evaluation Works", 2)

colored_box(
    "  EVALUATION PIPELINE\n"
    "\n"
    "  Test Data (20% of Spider)       Model Generates SQL       Compare\n"
    "\n"
    "  +-----------------------+       +-------------------+     +------------------+\n"
    "  | Question:             |       | Generated:        |     | Match?           |\n"
    "  | How many employees?   | ----> | SELECT COUNT(*)   | --> | Expected:        |\n"
    "  |                       |       | FROM employees;   |     | SELECT COUNT(*)  |\n"
    "  | Expected SQL:         |       |                   |     | FROM employees;  |\n"
    "  | SELECT COUNT(*)       |       +-------------------+     |                  |\n"
    "  | FROM employees;       |                                 | EXACT MATCH!     |\n"
    "  +-----------------------+                                 +------------------+\n"
    "\n"
    "  Results broken down by type:\n"
    "  +------------------+------------------+------------------+\n"
    "  | Simple Selects   | Filtered Queries | Aggregations     |\n"
    "  | (SELECT ... FROM)| (WHERE ...)      | (COUNT, AVG ...) |\n"
    "  | 91% accuracy     | 84% accuracy     | 87% accuracy     |\n"
    "  +------------------+------------------+------------------+",
    GREEN_BG,
)

doc.add_page_break()

# ═══════════════════════════════════════
# 3. HIGH-LEVEL DESIGN
# ═══════════════════════════════════════

section("3. High-Level Design (HLD)")

body("This section shows how all the pieces fit together at a high level.")

section("System Overview", 2)

colored_box(
    "                         THE BIG PICTURE\n"
    "\n"
    "    [ User's Browser ]                    [ Server ]\n"
    "    +----------------+                    +-----------------+\n"
    "    |                |   HTTP Request     |                 |\n"
    "    |  Next.js App   | -----------------> |  FastAPI App    |\n"
    "    |                |                    |                 |\n"
    "    |  - Landing     |   JSON Response    |  Checks:        |\n"
    "    |  - Dashboard   | <----------------- |  Is model ready?|\n"
    "    |  - Chat        |                    |                 |\n"
    "    +----------------+                    +-----+-----------+\n"
    "          |                                     |\n"
    "          v                              YES    |    NO\n"
    "    [Browser Storage]                     |     |\n"
    "     Query History                        v     v\n"
    "                                   +------+ +--------+\n"
    "                                   | LoRA | | Rule   |\n"
    "                                   | Model| | Based  |\n"
    "                                   +------+ +--------+\n"
    "                                        \\     /\n"
    "                                         v   v\n"
    "                                   SQL + Confidence\n"
    "                                      + Latency"
)

section("Three-Layer Architecture", 2)

colored_box(
    "   LAYER 1: FRONTEND             LAYER 2: BACKEND            LAYER 3: AI MODEL\n"
    "   (What users see)              (The brain)                 (The engine)\n"
    "\n"
    "   +------------------+          +------------------+        +------------------+\n"
    "   |                  |          |                  |        |                  |\n"
    "   |  Landing Page    |          |  Input Validator |        |  TinyLlama 1.1B  |\n"
    "   |                  |          |                  |        |                  |\n"
    "   |  Dashboard       |  ------> |  SQL Generator   | -----> |  LoRA Adapter    |\n"
    "   |   - Schema Box   |  HTTP    |                  | runs   |  (4.5 MB)        |\n"
    "   |   - Question Box |          |  Confidence Calc |        |                  |\n"
    "   |   - SQL Output   | <------  |                  | <----- |  Token Output    |\n"
    "   |                  |  JSON    |  Fallback Logic  |        |  + Probabilities |\n"
    "   |  Query History   |          |                  |        |                  |\n"
    "   +------------------+          +------------------+        +------------------+\n"
    "\n"
    "   Next.js + React               FastAPI + Python             HuggingFace + PyTorch\n"
    "   Port 3000                     Port 8000                    Loaded in memory",
    ACCENT_BG,
)

section("Simple Explanation", 2)
body("Think of it like a translator service:")
numbered("The user types what they want in English")
numbered("The frontend sends this to the backend server")
numbered("The server uses the AI model to \"translate\" English into SQL")
numbered("If the AI model is not available, a simpler rule-based system tries its best")
numbered("The result comes back with a confidence score (how sure the system is)")
numbered("The frontend displays the SQL, confidence, and how long it took")

section("Real World Comparison", 2)
body("This is similar to how these products work:")
nice_table(
    ["Product", "What They Do", "How This Project Compares"],
    [
        ["Google BigQuery NL", "Ask questions about data in English", "Same idea, but this uses a custom fine-tuned model"],
        ["GitHub Copilot", "Writes code from descriptions", "Similar approach (fine-tuned LLM) but for SQL instead of general code"],
        ["Salesforce Einstein", "Ask questions about CRM data", "Same concept applied to a different data domain"],
        ["AI2SQL / DuoSQL", "Standalone text-to-SQL tools", "This project shows you can build the core technology yourself"],
    ],
)

doc.add_page_break()

# ═══════════════════════════════════════
# 4. LOW-LEVEL DESIGN
# ═══════════════════════════════════════

section("4. Low-Level Design (LLD)")

section("4.1 What Happens When User Clicks 'Generate SQL'", 2)

body("Here is the step-by-step journey of a single request:")

colored_box(
    "  USER CLICKS GENERATE\n"
    "         |\n"
    "         v\n"
    "  [1] Frontend validates input\n"
    "      - Is schema empty? Show error toast\n"
    "      - Is question empty? Show error toast\n"
    "         |\n"
    "         v\n"
    "  [2] useGenerateSQL hook fires\n"
    "      - Sets loading = true\n"
    "      - Clears old results\n"
    "         |\n"
    "         v\n"
    "  [3] Axios sends POST /generate_sql\n"
    "      Body: { question, schema }\n"
    "         |\n"
    "         v\n"
    "  [4] FastAPI receives and validates\n"
    "      - Non-empty check\n"
    "      - Meaningful text check\n"
    "         |\n"
    "         v\n"
    "  [5] Is AI model loaded?\n"
    "      YES --> Build prompt, run model, get confidence\n"
    "      NO  --> Parse schema, match patterns, build SQL\n"
    "         |\n"
    "         v\n"
    "  [6] Return { sql, confidence, latency_ms }\n"
    "         |\n"
    "         v\n"
    "  [7] Frontend receives response\n"
    "      - Shows SQL in code block\n"
    "      - Shows confidence badge (green/amber/red)\n"
    "      - Shows latency\n"
    "      - Saves to history",
    GREEN_BG,
)

section("4.2 How Confidence Scoring Works", 2)

body("When the AI model generates SQL, it picks one word (token) at a time. For each token, the model has a probability score showing how confident it is about that choice.")

colored_box(
    "  HOW CONFIDENCE IS CALCULATED\n"
    "\n"
    "  Generated SQL:  SELECT  COUNT(*)  FROM  employees  WHERE  salary  >  50000\n"
    "  Token Probs:    0.95    0.91      0.97  0.88       0.93   0.86    0.94  0.82\n"
    "                   ^                                                       ^\n"
    "                 high                                                   lowest\n"
    "\n"
    "  Step 1: Average of all probs (geometric mean) = 0.907\n"
    "  Step 2: Find the weakest token                 = 0.82 (the number 50000)\n"
    "  Step 3: Blend them together\n"
    "          Final = (0.80 x 0.907) + (0.20 x 0.82)\n"
    "          Final = 0.726 + 0.164\n"
    "          Final = 0.89 = 89%  --> GREEN badge\n"
    "\n"
    "  Why this works: even if one word is uncertain, the score drops.\n"
    "  Users get a reliable trust signal.",
    ORANGE_BG,
)

body("The formula blends two things:")
bullet("80% weight on overall average quality across all tokens")
bullet("20% weight on the weakest token (the one the model was least sure about)")

body("This means: if the model was confident about every word, the score is high. If even one word was uncertain, the score drops.")

nice_table(
    ["Confidence Score", "Badge Color", "What It Means"],
    [
        ["80% or higher", "Green", "Model is quite confident. SQL is likely correct."],
        ["50% to 79%", "Amber/Yellow", "Model had some uncertainty. Review the SQL before running."],
        ["Below 50%", "Red", "Model struggled. Treat this as a rough suggestion only."],
    ],
)

section("4.3 API Request and Response Format", 2)

body("Here is exactly what the frontend sends and what the backend returns:")

colored_box(
    "  REQUEST (Frontend sends this)         RESPONSE (Backend returns this)\n"
    "  POST /generate_sql                    200 OK\n"
    "\n"
    "  {                                     {\n"
    "    \"question\": \"How many              \"sql\": \"SELECT COUNT(*)\n"
    "      employees earn                       FROM employees\n"
    "      more than 50000?\",                  WHERE salary > 50000;\",\n"
    "    \"schema\": \"CREATE TABLE            \"confidence\": 0.89,\n"
    "      employees (                        \"latency_ms\": 340,\n"
    "        id INT,                          \"model_used\": \"lora\"\n"
    "        name TEXT,                     }\n"
    "        salary INT\n"
    "      );\"\n"
    "  }\n"
    "\n"
    "  HEALTH CHECK                          HEALTH RESPONSE\n"
    "  GET /health                           200 OK\n"
    "\n"
    "                                        {\n"
    "                                          \"status\": \"healthy\",\n"
    "                                          \"model_loaded\": true\n"
    "                                        }",
    ACCENT_BG,
)

section("4.4 How the Fallback System Works", 2)

body("When the AI model is not loaded, the system still generates useful SQL by matching patterns in the question:")

nice_table(
    ["If the question contains...", "The system generates...", "Example"],
    [
        ["'how many'", "SELECT COUNT(*)", "How many employees? --> SELECT COUNT(*) FROM employees;"],
        ["'average' or 'avg'", "SELECT AVG(column)", "Average salary? --> SELECT AVG(salary) FROM employees;"],
        ["'maximum' or 'max'", "SELECT MAX(column)", "Max price? --> SELECT MAX(price) FROM products;"],
        ["'list' or 'show'", "SELECT columns FROM table", "Show all names --> SELECT name FROM employees;"],
        ["'more than' + number", "WHERE column > number", "Salary > 50000 --> WHERE salary > 50000"],
        ["'group by' or 'per'", "GROUP BY column", "Per department --> GROUP BY department"],
    ],
)

section("4.5 Frontend Components", 2)

nice_table(
    ["Component", "What It Does"],
    [
        ["Landing Page (/)", "Marketing page with features and call-to-action button"],
        ["Dashboard (/dashboard)", "Main tool: schema editor + question input + SQL output"],
        ["Chat (/chat)", "Chat-style input interface (demo, not connected to API yet)"],
        ["Header", "Shows project name, backend status (online/offline), theme toggle"],
        ["Schema Editor", "Text area where users paste CREATE TABLE statements"],
        ["Query Input", "Text area for the English question + Generate button"],
        ["SQL Output", "Displays generated SQL, confidence badge, latency, copy button"],
        ["Query History", "Sidebar showing past queries, click to reload any previous query"],
        ["Toast System", "Pop-up messages for success, errors, and info"],
    ],
)

colored_box(
    "  DASHBOARD PAGE LAYOUT\n"
    "\n"
    "  +================================================================+\n"
    "  |  HEADER: Logo    [Backend: Online]    [Dark/Light Toggle]      |\n"
    "  +================================================================+\n"
    "  |                                    |                           |\n"
    "  |  LEFT SIDE (Input)                 |  RIGHT SIDE (Output)      |\n"
    "  |                                    |                           |\n"
    "  |  +----------------------------+   |  +------------------------+|\n"
    "  |  | Schema Editor              |   |  | Generated SQL          ||\n"
    "  |  | CREATE TABLE employees ... |   |  | SELECT COUNT(*)        ||\n"
    "  |  |                            |   |  | FROM employees         ||\n"
    "  |  +----------------------------+   |  | WHERE salary > 50000;  ||\n"
    "  |                                    |  +------------------------+|\n"
    "  |  +----------------------------+   |                           |\n"
    "  |  | Question Input             |   |  Confidence: 92% [GREEN]  |\n"
    "  |  | How many employees earn    |   |  Latency: 340 ms          |\n"
    "  |  | more than 50000?           |   |  [Copy SQL]               |\n"
    "  |  +----------------------------+   |                           |\n"
    "  |                                    |  QUERY HISTORY            |\n"
    "  |  [ GENERATE SQL ]                 |  - How many employees...  |\n"
    "  |                                    |  - Show avg salary...     |\n"
    "  |                                    |  - List all products...   |\n"
    "  +================================================================+",
    ACCENT_BG,
)

section("4.6 How the Training Pipeline Works", 2)

body("LoRA (Low-Rank Adaptation) is a technique that lets you fine-tune a large model by only training a tiny fraction of its parameters.")

colored_box(
    "                     TRAINING PIPELINE\n"
    "\n"
    "  +---------------+     +-----------------+     +-------------------+\n"
    "  |  Spider CSV   |     |  Data Pipeline  |     |  Instruction      |\n"
    "  |  8,035 rows   | --> |  (pipeline.py)  | --> |  Prompts          |\n"
    "  |               |     |                 |     |                   |\n"
    "  |  question     |     |  - Read CSV     |     |  SYSTEM: You are  |\n"
    "  |  sql          |     |  - Format       |     |  a SQL expert...  |\n"
    "  |  schema       |     |  - Split 80/20  |     |  USER: schema +   |\n"
    "  +---------------+     +-----------------+     |  question         |\n"
    "                                                |  ASSISTANT: sql   |\n"
    "                                                +-------------------+\n"
    "                                                         |\n"
    "                                                         v\n"
    "  +---------------+     +-----------------+     +-------------------+\n"
    "  | Saved Adapter |     |   SFT Trainer   |     |  TinyLlama +      |\n"
    "  | (4.5 MB)      | <-- |   3 epochs      | <-- |  LoRA Adapters    |\n"
    "  |               |     |   lr=0.0002     |     |  (rank 8)         |\n"
    "  | Ready for     |     |   cosine decay  |     |                   |\n"
    "  | inference!    |     |   checkpoints   |     |  99.9% frozen     |\n"
    "  +---------------+     +-----------------+     |  0.1% trainable   |\n"
    "                                                +-------------------+",
    GREEN_BG,
)

section("LoRA Training Configuration", 3)

nice_table(
    ["Setting", "Value", "Why"],
    [
        ["Base Model", "TinyLlama 1.1B Chat", "Small enough for laptop training, already good at following instructions"],
        ["LoRA Rank", "8", "Good balance of quality and efficiency"],
        ["LoRA Alpha", "16", "Standard 2x multiplier of rank"],
        ["Dropout", "0.05", "Light regularization to prevent overfitting"],
        ["Target Layers", "q_proj, v_proj", "Attention query and value layers are most impactful for generation"],
        ["Trainable Params", "1.1M (0.10%)", "99.9% of model stays frozen"],
        ["Adapter File Size", "4.5 MB", "Compare to 2.2 GB for the full model"],
        ["Training Epochs", "3", "Three passes over all 8,035 examples"],
        ["Learning Rate", "0.0002", "With cosine schedule and warmup"],
        ["Batch Size", "1 x 16 accumulation = 16 effective", "Small batch with gradient accumulation for memory savings"],
    ],
)

section("LoRA vs Full Fine-Tuning: Visual Comparison", 3)

colored_box(
    "  FULL FINE-TUNING                     LoRA FINE-TUNING (What I Used)\n"
    "\n"
    "  +========================+           +========================+\n"
    "  | Layer 1   [UPDATING]   |           | Layer 1   [FROZEN]     |\n"
    "  | Layer 2   [UPDATING]   |           | Layer 2   [FROZEN]     |\n"
    "  | Layer 3   [UPDATING]   |           | Layer 3   [FROZEN]     |\n"
    "  | ...                    |           | ...                    |\n"
    "  | Layer 22  [UPDATING]   |           | Layer 22  [FROZEN]     |\n"
    "  +========================+           +======+====+============+\n"
    "                                              |    |\n"
    "  All 1.1 BILLION parameters                  v    v\n"
    "  updated every step                     +----+----+----+\n"
    "                                         | LoRA Adapter |\n"
    "  Needs: 24 GB+ GPU memory               | 1.1 MILLION  |\n"
    "  Takes: Days of training                | params only  |\n"
    "  Output: 2.2 GB model file              +-------------+\n"
    "                                         Needs: 8 GB RAM\n"
    "                                         Takes: 2.5 hours\n"
    "                                         Output: 4.5 MB file",
    ORANGE_BG,
)

doc.add_page_break()

# ═══════════════════════════════════════
# 5. DEMO SCRIPT
# ═══════════════════════════════════════

section("5. Live Demo Script (8-10 Minutes)")

section("Before the demo", 2)
body("Run this command to start everything: ./run.sh")
body("Or start manually: python api/serve.py + cd frontend && npx next dev")

section("Minute 1-2: Introduction", 2)
body("\"I built an end-to-end system that converts English questions into SQL queries. Let me show you how it works.\"")
body("Open the landing page (localhost:3000). Point out the features and click 'Open Dashboard'.")

section("Minute 3-6: Core Demo", 2)
numbered("Point to the green 'Backend Connected' indicator in the header")
numbered("Select the 'Employees' sample schema from the dropdown")
numbered("Type: \"How many employees earn more than 50000?\"")
numbered("Click Generate. While loading, say: \"This sends the question to our FastAPI backend which runs the LoRA model\"")
numbered("When result shows, point out: the SQL query, the green confidence badge, the latency, the copy button")
numbered("Show the history sidebar with the query saved")
numbered("Try another: \"Show average salary by department\" to show a different query type")

section("Minute 7-8: Architecture Explanation", 2)
body("\"The frontend sends a POST request to the backend. The backend builds a prompt with the schema and question, runs it through our fine-tuned model, calculates a confidence score from token probabilities, and returns structured JSON. If the model is not available, it automatically falls back to rule-based SQL generation.\"")

section("Minute 9-10: Technical Depth", 2)
body("\"I trained on 8,000 Spider examples using LoRA, which means only 0.1% of the model parameters are trainable. The adapter is just 4.5 MB. The confidence score blends average token probability with minimum token probability, so even one uncertain word brings the score down. This gives users a trustworthy reliability signal.\"")

doc.add_page_break()

# ═══════════════════════════════════════
# 6. INTERVIEW PITCHES
# ═══════════════════════════════════════

section("6. Interview Pitches")

section("30-Second Pitch", 2)
colored_box(
    "I built a full-stack Text-to-SQL system. Users paste their database\n"
    "schema and ask questions in English. The backend runs a LoRA fine-tuned\n"
    "model that generates SQL with confidence scores. The frontend is a\n"
    "Next.js dashboard with schema editing, health monitoring, and query\n"
    "history. If the model is offline, the system falls back to rule-based\n"
    "SQL so it never breaks.",
    ACCENT_BG,
)

section("2-Minute Pitch", 2)
body("The system has three layers:")
numbered("ML Pipeline: I fine-tuned TinyLlama on 8,000 Spider examples using LoRA. Only 0.1% of parameters are trainable, producing a 4.5 MB adapter. Training takes 2.5 hours on a MacBook.")
numbered("API Layer: FastAPI exposes /generate_sql and /health. When the model is loaded, it generates SQL with confidence from token probabilities. When it is not loaded, smart pattern matching provides fallback SQL.")
numbered("Frontend: Next.js app with landing page, dashboard (schema + question + SQL output), health checks, and persistent query history.")

section("5-Minute Deep Dive", 2)
body("Start with the problem: SQL is powerful but hard for non-technical people. They know what data they need but cannot write the query.")
body("My approach: supervised fine-tuning with LoRA on a lightweight model. LoRA freezes 99.9% of the model and only trains small adapter matrices in the attention layers. This makes training fast and the adapter tiny.")
body("The backend handles reliability: input validation catches bad requests, confidence scoring tells users how trustworthy the result is, and fallback mode keeps the system available even during model failures.")
body("The frontend follows production patterns: custom hooks for API integration, toast notifications for feedback, localStorage history for continuity, and health monitoring for status awareness.")
body("Results: 0.40 training loss, 87.5% token accuracy, and sub-second inference. The evaluation breaks down accuracy by query type so I can see exactly where the model needs improvement.")

doc.add_page_break()

# ═══════════════════════════════════════
# 7. DESIGN DECISIONS
# ═══════════════════════════════════════

section("7. Design Decisions and Trade-offs")

nice_table(
    ["Decision", "Why I Chose This", "Alternative", "Trade-off"],
    [
        ["TinyLlama 1.1B", "Fits on laptop, fast training and inference", "Bigger model (7B, 13B)", "Less capacity for complex queries"],
        ["LoRA fine-tuning", "0.1% params, 4.5 MB adapter, fast to iterate", "Full fine-tuning", "Slightly lower ceiling on complex SQL"],
        ["FastAPI backend", "Async, auto-docs, Pydantic validation, lightweight", "Flask or Django", "Smaller ecosystem for auth/ORM"],
        ["Rule-based fallback", "System stays useful when model is down", "Just return error", "Fallback quality is lower"],
        ["Token-prob confidence", "Meaningful uncertainty signal from model internals", "No confidence", "Small latency cost for output_scores"],
        ["localStorage history", "Simple, works offline, no backend state needed", "Database storage", "Not portable across devices"],
        ["Spider dataset", "8K high-quality examples, diverse patterns", "WikiSQL (80K simpler)", "Smaller but higher quality per example"],
        ["Next.js frontend", "File routing, industry standard, React ecosystem", "Vite + React Router", "Heavier framework"],
    ],
    col_widths=[3.5, 5, 3.5, 4],
)

doc.add_page_break()

# ═══════════════════════════════════════
# 8. AI QUESTIONS
# ═══════════════════════════════════════

section("8. How AI Helps and AI-Related Interview Questions")

section("How AI is used in this project", 2)
bullet("The core feature IS AI: a fine-tuned model converts English to SQL", "Main Feature")
bullet("Confidence scoring uses the model's own token probabilities to create a trust signal", "Reliability")
bullet("LoRA itself is an AI research technique that makes fine-tuning practical on consumer hardware", "Efficiency")

section("How AI could extend this project further", 2)
bullet("Use AI to explain generated SQL in plain English so users understand what the query does", "SQL Explanation")
bullet("Use AI to auto-suggest questions based on the database schema", "Smart Suggestions")
bullet("Use embeddings to find similar past queries and show cached results for faster response", "Semantic Caching")
bullet("Use AI to validate generated SQL by checking if it makes sense with the schema", "AI Validation")
bullet("Fine-tune on user feedback (thumbs up/down) to continuously improve over time", "Self-Improvement")
bullet("Use RAG (Retrieval-Augmented Generation) to pull similar examples at inference time", "RAG Enhancement")
bullet("Stream SQL tokens to the frontend in real-time so users see results appearing", "Streaming Output")

section("AI Interview Questions", 2)

qa("Why did you fine-tune instead of using GPT-4 or Claude API?",
   "Three reasons. First, cost: API calls for every query get expensive at scale. Second, latency: local inference is faster than round-tripping to an external API. Third, control: I own the model, can run it offline, and do not depend on third-party pricing or availability changes. LoRA makes fine-tuning practical even on a laptop.")

qa("How do you handle hallucination (model generating wrong SQL)?",
   "Multiple layers. The confidence score flags uncertain outputs. Low confidence means the model was not sure. The fallback system uses deterministic rules that cannot hallucinate. In production, I would add SQL parsing to verify the output is valid SQL and check that table/column names match the actual schema before returning results.")

qa("What is LoRA and why is it better than full fine-tuning?",
   "LoRA freezes the entire pretrained model and adds small trainable matrices into attention layers. Instead of updating 1.1 billion parameters, I update only 1.1 million. The result is a 4.5 MB adapter file instead of a 2.2 GB model. This means faster training, lower memory usage, cheaper deployment, and I can even swap different adapters on the same base model for different tasks.")

qa("How would you use AI to improve the frontend experience?",
   "Several ways: auto-suggest questions based on schema, show plain-English explanations next to generated SQL, stream tokens in real-time so users see SQL appearing word by word, and add a 'fix this query' button that sends errors back to the model for correction.")

qa("How reliable is the confidence score?",
   "It is not a guarantee of correctness, but it correlates well with actual accuracy. High-confidence queries tend to be correct because the model was sure about every token. The score combines average quality (geometric mean of all token probs) with worst-case quality (minimum token prob), so even one uncertain word pulls the score down. In the UI, color coding (green/amber/red) helps users quickly decide whether to trust the result.")

doc.add_page_break()

# ═══════════════════════════════════════
# 9. REAL-WORLD USE CASES
# ═══════════════════════════════════════

section("9. Real-World Use Cases")

section("Who would use this?", 2)
nice_table(
    ["User", "Situation", "Example Question"],
    [
        ["Business Analyst", "Needs quick metrics without writing SQL", "Show total revenue by region for Q4"],
        ["Product Manager", "Wants to check data to validate a hypothesis", "How many users signed up last month?"],
        ["Customer Support", "Looks up info during customer calls", "Find all orders for customer John Smith"],
        ["Data Science Student", "Learning SQL by seeing translations", "What is the average GPA of CS majors?"],
        ["Non-Technical Executive", "Needs insights from raw data", "Which department has the most employees?"],
    ],
)

section("What a production version would look like", 2)

colored_box(
    "              PRODUCTION ARCHITECTURE\n"
    "\n"
    "    [ Users ]     [ Users ]     [ Users ]\n"
    "        \\            |            /\n"
    "         v           v           v\n"
    "    +-----------------------------------+\n"
    "    |         Load Balancer              |\n"
    "    +-----------------------------------+\n"
    "              |          |\n"
    "         +----+----+  +-+------+\n"
    "         | API     |  | API    |\n"
    "         | Server 1|  | Server 2|\n"
    "         +----+----+  +---+----+\n"
    "              |           |\n"
    "         +----+-----------+----+\n"
    "         |  Model Server       |\n"
    "         |  (vLLM on GPU)      |\n"
    "         +---------------------+\n"
    "         |  Monitoring         |\n"
    "         |  (Logs + Metrics)   |\n"
    "         +---------------------+",
    ORANGE_BG,
)

doc.add_page_break()

# ═══════════════════════════════════════
# 10. SCALING ROADMAP
# ═══════════════════════════════════════

section("10. Scaling and Production Roadmap")

section("Phase 1: Hardening (1-2 weeks)", 2)
bullet("Add request logging with timestamps, query text, confidence, and latency")
bullet("Add backend API tests for validation and response format")
bullet("Connect /chat page to the real /generate_sql endpoint")
bullet("Add SQL syntax checking as a post-processing step")

section("Phase 2: Quality and Safety (2-4 weeks)", 2)
bullet("Add SQL execution sandbox with strict timeouts and row limits")
bullet("Improve training data with schema context in prompts")
bullet("Add experiment tracking (Weights & Biases or MLflow)")
bullet("Add streaming response support for real-time token display")

section("Phase 3: Scale (4-8 weeks)", 2)
bullet("Move to vLLM or TGI for high-throughput batched inference")
bullet("Add user authentication and rate limiting")
bullet("Deploy behind Nginx/Caddy with HTTPS")
bullet("Add monitoring dashboard (Prometheus + Grafana)")
bullet("Set up CI/CD pipeline with smoke tests")

doc.add_page_break()

# ═══════════════════════════════════════
# 11. FULL Q&A BANK
# ═══════════════════════════════════════

section("11. Interview Q&A Bank (20+ Questions)")

qa("Tell me about this project.",
   "I built an end-to-end Text-to-SQL platform. Users paste a database schema and ask questions in English. The backend generates SQL using a LoRA fine-tuned model and returns confidence scores and latency. The frontend is a Next.js dashboard with schema editing, real-time health monitoring, and query history. The system also has a fallback mode that generates rule-based SQL when the model is not available.")

qa("What was the hardest part?",
   "Balancing model quality with hardware constraints. Training a language model usually needs expensive GPUs, but LoRA let me fine-tune on a MacBook by training only 0.1% of parameters. The second challenge was making confidence scoring meaningful. I tested several methods before choosing the geometric-mean plus minimum-probability blend.")

qa("How does the model work?",
   "TinyLlama is a pretrained language model that generates text token by token. I fine-tuned it with LoRA, which adds small trainable matrices to the attention layers while keeping everything else frozen. During training, the model learns from 8,000 examples of English questions paired with correct SQL. At inference, I give it the schema and question in a structured prompt, and it generates SQL one token at a time.")

qa("Why not just use ChatGPT?",
   "Cost, latency, and control. ChatGPT API calls cost money for every query and add network latency. With fine-tuning, I own the model, it runs locally with no internet needed, and I do not depend on OpenAI's pricing or availability. Plus, a specialized fine-tuned model often outperforms a general one on focused tasks.")

qa("How do you know the SQL is correct?",
   "Multiple signals. The confidence score from token probabilities flags uncertain generations. The evaluation harness checks exact match accuracy against known-correct SQL, broken down by query type (simple, filter, aggregation). In production, I would add SQL parsing validation and sandboxed execution to test queries before returning them.")

qa("How would you deploy this?",
   "Separate the model server (vLLM on GPU) from the API layer (stateless FastAPI replicas behind a load balancer). Add auth, rate limiting, logging, and monitoring. Frontend goes to Vercel or a CDN. Use CI/CD with automated tests on each deploy.")

qa("What would you improve with more time?",
   "Four things: add schema context to training prompts for better accuracy, add streaming SQL generation for better user experience, build a SQL execution sandbox so users can run queries directly, and add RLHF-based learning from user feedback.")

qa("What is the confidence score?",
   "When the model generates each SQL token, it has a probability for each possible word. I take the probability of the word it actually chose. The final confidence is 80% geometric mean of all token probabilities (overall quality) plus 20% minimum token probability (weakest point). High score means the model was sure about every word.")

qa("What is LoRA?",
   "Low-Rank Adaptation. Instead of updating all 1.1 billion model parameters, LoRA adds small matrices (rank 8) to specific attention layers and trains only those. This reduces trainable parameters by 1000x, makes training possible on a laptop, and produces a 4.5 MB adapter instead of a 2.2 GB model file.")

qa("Why TinyLlama?",
   "It fits in memory on a laptop, trains in hours instead of days, and serves inference in under a second. For a focused task like text-to-SQL, a smaller fine-tuned model often beats a larger general-purpose model because all its capacity is directed at the target task.")

qa("How does the frontend talk to the backend?",
   "The frontend uses Axios to send HTTP requests to the backend URL (set via environment variable). A custom React hook called useGenerateSQL manages the full lifecycle: sets loading state, sends the POST request, handles the response or error, and updates the UI state. A separate hook polls /health every 15 seconds for status.")

qa("What happens if the backend goes down?",
   "The health check detects the outage within 15 seconds and shows a red 'Backend Offline' indicator. The Generate button gets disabled. If the model specifically fails but the API server is still running, the backend automatically switches to rule-based SQL generation. The app stays useful in both cases.")

qa("How did you evaluate the model?",
   "I run the evaluation script on the test split. It generates SQL for each example, normalizes both predicted and expected SQL (lowercase, whitespace cleanup), and checks for exact match. Results are broken down by query category: simple selects, filtered queries, and aggregations.")

qa("What datasets did you use?",
   "Primarily Spider CSV with 8,035 high-quality text-to-SQL examples covering diverse patterns. The system also supports WikiSQL (80K simpler examples) and SQL-Create-Context (78K with schema). I chose Spider for quality and variety.")

qa("How would you add multi-database support?",
   "Add a schema registry that stores metadata for multiple databases. The frontend lets users pick which database to query. The backend fetches the right schema before building the prompt. Training data would need schema-enriched prompts.")

qa("How would you handle sensitive data?",
   "Never expose raw data through results. Add column-level access controls. Log every generated query for auditing. Use parameterized execution to prevent injection. The model only sees schema structure, never actual data.")

qa("What testing would you add?",
   "Backend: pytest for validation, mock inference, fallback logic, and response format. Frontend: React Testing Library for hook behavior and component rendering. Integration: Playwright end-to-end tests for the full generate flow.")

qa("What is SFT vs RLHF?",
   "SFT (Supervised Fine-Tuning) trains directly on input-output pairs. The model learns to produce correct SQL given an instruction. RLHF adds human preference feedback and reinforcement learning on top. SFT is simpler and sufficient when you have high-quality SQL targets. RLHF would help later for handling ambiguous queries where user preference matters.")

qa("How would you make inference faster?",
   "Use vLLM for continuous batching (handles multiple requests efficiently). Use quantization (4-bit or 8-bit) to reduce memory and speed up math. Add response caching for repeated queries. Use KV-cache optimization to avoid recomputing attention for the prompt.")

qa("Explain the training configuration.",
   "3 epochs over 8K examples. Effective batch size 16 (batch 1 with 16 gradient accumulation steps to save memory). Cosine learning rate starting at 0.0002 with 5% warmup. Float16 precision for speed. Gradient checkpointing enabled for memory efficiency. Checkpoints every 200 steps with best model saved based on lowest eval loss.")

doc.add_page_break()

# ═══════════════════════════════════════
# 12. RESUME BULLETS
# ═══════════════════════════════════════

section("12. Resume-Ready Bullets")

bullet("Architected and delivered an end-to-end Text-to-SQL platform with LoRA fine-tuning, FastAPI inference, and Next.js dashboard, achieving 87.5% token accuracy and sub-second latency.")
bullet("Implemented confidence-scored SQL generation using token-level log probabilities, surfacing model reliability metrics directly in the product interface with color-coded badges.")
bullet("Designed dual-mode inference with LoRA model path and deterministic rule-based fallback, ensuring 100% API availability regardless of model state.")
bullet("Built modular React frontend with custom hooks for API integration, health monitoring, and persistent query history with localStorage.")
bullet("Fine-tuned TinyLlama 1.1B using LoRA (r=8, 0.10% trainable params, 4.5 MB adapter) on 8K Spider examples with evaluation breakdown by query category.")
bullet("Created reproducible ML pipeline with YAML configuration, CLI scripts, and one-command bootstrap for local development.")

doc.add_page_break()

# ═══════════════════════════════════════
# 13. CODE WALKTHROUGH ORDER
# ═══════════════════════════════════════

section("13. Code Walkthrough Order (For Interview)")

body("If asked to walk through the code, follow this order. It tells the story from what the user sees to how it works underneath:")

nice_table(
    ["Order", "File", "What To Explain"],
    [
        ["1", "frontend/app/dashboard/page.tsx", "User journey: schema input, question, generate button, result display"],
        ["2", "frontend/hooks/useGenerateSQL.ts", "How the frontend calls the API: loading, error, result states"],
        ["3", "frontend/lib/api.ts", "The HTTP client: Axios setup, POST request, error handling"],
        ["4", "api/serve.py", "Backend logic: validation, model vs fallback branch, confidence, response"],
        ["5", "scripts/train_spider.py", "Training story: data loading, LoRA setup, SFTTrainer, saving"],
        ["6", "data/pipeline.py", "How training data is prepared: CSV reading, prompt formatting, splitting"],
        ["7", "evaluation/evaluator.py", "Quality measurement: exact match, normalization, category breakdown"],
    ],
)

body("This sequence flows from product to integration to backend to ML pipeline to quality assurance, which matches how interviewers think about systems.")

# ═══════════════════════════════════════
# SAVE
# ═══════════════════════════════════════

output_path = "Text-to-SQL_Interview_Documentation.docx"
doc.save(output_path)
print(f"Done! Saved to: {output_path}")
